from __future__ import annotations

import datetime as _dt
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_PATH = ROOT / "docs" / "System_Zabezpieczen_HA_Z2M_NodeRED_Viktor-Art_PL.docx"


PALETTE = {
    "primary": RGBColor(0x0D, 0x2D, 0x4F),  # deep navy
    "accent": RGBColor(0x1E, 0x88, 0xE5),  # blue
    "ok": RGBColor(0x2E, 0x7D, 0x32),  # green
    "warn": RGBColor(0xEF, 0x6C, 0x00),  # orange
    "alarm": RGBColor(0xC6, 0x28, 0x28),  # red
    "muted": RGBColor(0x54, 0x60, 0x6A),  # gray
}


def _set_run_style(run, *, bold=False, color=None, size_pt: int | None = None):
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    if size_pt is not None:
        run.font.size = Pt(size_pt)


def _set_cell_shading(cell, hex_color: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _add_hr(paragraph, color_hex: str = "D9DDE3"):
    p = paragraph._p
    p_pr = p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "10")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color_hex)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def _h1(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    _set_run_style(r, bold=True, color=PALETTE["primary"], size_pt=22)
    _add_hr(doc.add_paragraph())


def _h2(doc: Document, text: str):
    p = doc.add_paragraph()
    r = p.add_run(text)
    _set_run_style(r, bold=True, color=PALETTE["primary"], size_pt=16)


def _note(doc: Document, title: str, body: str, kind: str):
    # kind: ok/warn/alarm/info
    color = {
        "ok": ("E8F5E9", PALETTE["ok"]),
        "warn": ("FFF3E0", PALETTE["warn"]),
        "alarm": ("FFEBEE", PALETTE["alarm"]),
        "info": ("E3F2FD", PALETTE["accent"]),
    }[kind]
    fill, line = color

    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.rows[0].cells[0]
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    _set_cell_shading(cell, fill)

    p = cell.paragraphs[0]
    r = p.add_run(title + "  ")
    _set_run_style(r, bold=True, color=line, size_pt=11)
    r2 = p.add_run(body)
    _set_run_style(r2, color=PALETTE["primary"], size_pt=11)

    doc.add_paragraph()


def _bullet(doc: Document, text: str):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.space_after = Pt(2)


def _code(doc: Document, text: str):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(10)
    run.font.color.rgb = PALETTE["muted"]


def main():
    now = _dt.datetime.now().astimezone()
    doc = Document()

    # Page setup: A4 portrait, comfortable margins
    section = doc.sections[0]
    section.page_height = Inches(11.69)  # A4
    section.page_width = Inches(8.27)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)

    # Cover
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("SYSTEM ZABEZPIECZEŃ\nHome Assistant + Zigbee2MQTT + Node-RED")
    _set_run_style(r, bold=True, color=PALETTE["primary"], size_pt=26)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Viktor‑Art — instrukcja premium dla użytkowników")
    _set_run_style(r, bold=False, color=PALETTE["muted"], size_pt=12)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = meta.add_run(f"Wersja: 1.0  •  Data: {now.strftime('%Y-%m-%d %H:%M')}  •  Obszar: dom + warsztat (≈2000 m²)")
    _set_run_style(r, color=PALETTE["muted"], size_pt=10)

    doc.add_paragraph()
    _note(
        doc,
        "Uwaga (bezpieczeństwo):",
        "Automatyka HA/Zigbee nie zastępuje certyfikowanych systemów przeciwpożarowych. "
        "To warstwa informacyjno‑automatyzacyjna zwiększająca komfort i szybkość reakcji.",
        "warn",
    )

    doc.add_page_break()

    _h1(doc, "1. Cel systemu i filozofia działania")
    doc.add_paragraph(
        "Celem jest szybkie wykrywanie zagrożeń (zalanie / dym / wysoka temperatura), "
        "automatyczne działania minimalizujące szkody (odcięcie zasilania pomp) oraz "
        "wielokanałowe powiadomienia i sygnalizacja, z możliwością potwierdzenia (ACK) i wyciszenia."
    )

    _h2(doc, "Warstwy ochrony (model 4‑warstwowy)")
    _bullet(doc, "Warstwa 0: urządzenia lokalne (syreny/detektory) — działają nawet bez HA.")
    _bullet(doc, "Warstwa 1: Zigbee2MQTT — zbiera zdarzenia z wielu niezależnych sieci Zigbee.")
    _bullet(doc, "Warstwa 2: Node‑RED — logika alarmów, eskalacja, wzorce syren, sterowanie pompami.")
    _bullet(doc, "Warstwa 3: Home Assistant — powiadomienia push/email + nadzór (supervision) i dashboard.")
    _bullet(doc, "Dodatkowo: przypomnienia push eskalują razem z alarmem (email nie eskaluje), ale są dławione minimalnym odstępem aby uniknąć spamu.")

    _note(
        doc,
        "Dlaczego to jest 'robust'?",
        "Masz separację sieci Zigbee (brak mieszania koordynatorów), nadzór nad dostępnością, "
        "test‑mode, manualne potwierdzenia oraz zabezpieczenia przed przypadkowym przywróceniem pomp.",
        "ok",
    )

    _h1(doc, "2. Architektura techniczna (skrót)")
    doc.add_paragraph("System składa się z następujących elementów:")

    table = doc.add_table(rows=1, cols=3)
    hdr = table.rows[0].cells
    hdr[0].text = "Element"
    hdr[1].text = "Rola"
    hdr[2].text = "Co się stanie, gdy padnie?"
    for c in hdr:
        _set_cell_shading(c, "E3F2FD")
        for run in c.paragraphs[0].runs:
            _set_run_style(run, bold=True, color=PALETTE["primary"], size_pt=10)

    rows = [
        ("Mosquitto (MQTT)", "Wspólna magistrala komunikatów", "Brak zdarzeń do Node‑RED/HA → brak automatyki."),
        ("Zigbee2MQTT (wiele instancji)", "Oddzielne sieci Zigbee (koordynatory)", "Dana sieć traci telemetrię/sterowanie."),
        ("Node‑RED", "Logika alarmów + sterowania", "Brak automatycznych reakcji, ale urządzenia lokalne nadal działają."),
        ("Home Assistant", "Push/email + dashboard + supervision", "Brak powiadomień i nadzoru, Node‑RED może nadal działać."),
    ]
    for a, b, c in rows:
        r = table.add_row().cells
        r[0].text = a
        r[1].text = b
        r[2].text = c

    doc.add_paragraph()

    _h1(doc, "3. Zigbee2MQTT — sieci i zabezpieczenia integralności")
    _h2(doc, "Najważniejsze zasady")
    _bullet(doc, "Każda instancja Zigbee2MQTT to osobna sieć Zigbee (unikalne: kanał + PAN + ExtPAN + network key).")
    _bullet(doc, "`permit_join` domyślnie WYŁĄCZONE; włączaj tylko na czas parowania.")
    _bullet(doc, "Frontend port może się powtarzać (Ingress HA), ale MQTT `base_topic` i `serial.port` muszą być unikalne.")
    _bullet(doc, "Gdy widzisz błąd `configuration-adapter mismatch` — nie startuj na siłę; napraw konfigurację/backup.")

    _h2(doc, "Supervision (nadzór) po stronie HA")
    _bullet(doc, "Bridge offline → push + email.")
    _bullet(doc, "Pairing (permit_join) zostawione ON > 15 min → push + email.")
    _bullet(doc, "Offline devices > 0 przez 10 min → push + email (wskazuje problemy z zasięgiem, routerami, bateriami).")

    _h1(doc, "4. Node‑RED — logika alarmów (co dokładnie się dzieje)")
    _h2(doc, "4.1 Typy alarmów")
    _bullet(doc, "Zalanie (FLOOD): czujnik `water_leak: true` → powiadomienia + (opcjonalnie) syreny + WYŁĄCZENIE pomp.")
    _bullet(doc, "Dym (SMOKE): `smoke/alarm: true` → powiadomienia + syreny w sekwencji (2 beepy warsztat / 3 beepy dom).")
    _bullet(doc, "Temperatura (HEAT): `temperature >= 45°C` → powiadomienia + syreny (osobny wzorzec).")

    _h2(doc, "4.2 Tryby pracy syren (Production/Test)")
    _bullet(doc, "Tryb testowy: syreny OFF, ale powiadomienia i logika nadal działają.")
    _bullet(doc, "Tryb produkcyjny: syreny ON zgodnie z wzorcami.")
    _bullet(doc, "Przełączanie: z dashboardu HA (przyciski) lub w Node‑RED (inject).")

    _h2(doc, "4.3 Potwierdzanie / wyciszanie (ACK)")
    doc.add_paragraph(
        "W systemie są przyciski fizyczne B (Zigbee) oraz przyciski wirtualne w HA. "
        "Uwaga: wyciszenie nie 'wyłącza' czujnika — ono zatrzymuje WYJŚCIA systemu (syreny / pętle)."
    )
    _bullet(doc, "1× naciśnięcie: wycisza i tłumi beepy na `silenceMinutes` (domyślnie 60 min).")
    _bullet(doc, "3× naciśnięcie w 5 s: reset pętli powtarzania + próba przywrócenia pomp (warunki niżej).")
    _bullet(doc, "Anty‑spam: pojedyncze naciśnięcie bez aktywnego alarmu nie wysyła powiadomień.")

    _h2(doc, "4.4 Pompy — wyłączenie i bezpieczne przywracanie")
    _note(
        doc,
        "Bezpieczny domyślny stan:",
        "Po wykryciu zalania pompy są wyłączane i pozostają wyłączone do czasu świadomego przywrócenia.",
        "info",
    )
    _bullet(doc, "Wyłączenie pomp: natychmiast przy FLOOD (na gniazdach P1/P2 przez MQTT Z2M).")
    _bullet(doc, "Przywrócenie pomp (ON): tylko po 3× naciśnięciu i spełnieniu warunków bezpieczeństwa.")

    _h2(doc, "Warunki przywrócenia pomp (3× w 5 s)")
    _bullet(doc, "Brak zdarzeń 'wet' przez ostatnie `pumpRestoreDryMinutes` (domyślnie 5 min).")
    _bullet(doc, "Każdy czujnik zalania musi wysłać raport DRY w ostatnich `pumpRestoreDryReportMinutes` (domyślnie 10 min).")
    _bullet(doc, "Raport DRY musi być późniejszy niż ostatni raport WET dla danego czujnika.")
    _bullet(doc, "Jeżeli system nie ma historii zalania (brak wcześniejszego WET) → nie przywraca pomp (nic do przywrócenia).")

    _h1(doc, "5. Home Assistant — dashboard i obsługa operatora")
    _h2(doc, "5.1 Panel 'Safety' na dashboardzie Zigbee2MQTT")
    _bullet(doc, "`sensor.safety_health`: OK/WARN/ALARM (szybki status całej instalacji).")
    _bullet(doc, "`sensor.safety_z2m_offline_devices_total`: suma offline urządzeń we wszystkich sieciach.")
    _bullet(doc, "`binary_sensor.safety_z2m_any_bridge_offline`: czy dowolny bridge Z2M jest offline.")
    _bullet(doc, "`binary_sensor.safety_z2m_any_permit_join_on`: czy gdziekolwiek jest włączone parowanie.")
    _bullet(doc, "`input_boolean.safety_sirens_enabled`: kontrolka trybu syren (test/production).")

    _h2(doc, "5.2 Przyciski operatora")
    _bullet(doc, "Silence alarms (virtual) — działa jak B‑przycisk 1×.")
    _bullet(doc, "Test mode (sirens off) — wyłącza syreny (powiadomienia zostają).")
    _bullet(doc, "Production mode (sirens on) — włącza syreny.")

    _h1(doc, "6. Procedury (Quick Start)")
    _h2(doc, "6.1 Testy bez ryzyka")
    _bullet(doc, "Włącz Test mode (sirens off).")
    _bullet(doc, "W Node‑RED użyj: TEST: Flood / Smoke / Heat oraz TEST: Flood (pumps off) jeśli chcesz sprawdzić pompę.")
    _bullet(doc, "Sprawdź: push + email; następnie użyj Silence 1× oraz 3× (reset).")

    _h2(doc, "6.2 Parowanie nowych urządzeń Zigbee")
    _bullet(doc, "Włącz permit_join tylko w tej instancji Z2M, która ma przyjąć urządzenie.")
    _bullet(doc, "Po sparowaniu natychmiast wyłącz permit_join (system supervision i tak ostrzeże po 15 min).")

    _h2(doc, "6.3 Co robić przy realnym zalaniu")
    _bullet(doc, "Sprawdź miejsce fizycznie, usuń przyczynę.")
    _bullet(doc, "Poczekaj na DRY raporty czujników (lub wymuś odświeżenie wg instrukcji czujnika).")
    _bullet(doc, "Dopiero wtedy użyj 3× (w 5s) aby przywrócić pompy.")

    _h1(doc, "7. Troubleshooting (najczęstsze przypadki)")
    _bullet(doc, "Brak powiadomień: sprawdź Node‑RED Debug: `DEBUG HA response` + tokeny/URL.")
    _bullet(doc, "Brak syren: upewnij się, że tryb Production jest włączony (kontrolka w HA).")
    _bullet(doc, "Pompy nie wracają ON: sprawdź warunki DRY/WET (ostatnie raporty) i okno czasowe 5/10 minut.")
    _bullet(doc, "Dużo offline: dodaj routery Zigbee, sprawdź baterie, interferencje Wi‑Fi, rozmieszczenie koordynatorów.")

    doc.add_page_break()
    _h1(doc, "Załącznik A — najważniejsze parametry (domyślne)")
    t = doc.add_table(rows=1, cols=3)
    h = t.rows[0].cells
    h[0].text = "Parametr"
    h[1].text = "Wartość"
    h[2].text = "Znaczenie"
    for c in h:
        _set_cell_shading(c, "E8EAF6")
        for run in c.paragraphs[0].runs:
            _set_run_style(run, bold=True, color=PALETTE["primary"], size_pt=10)

    for k, v, d in [
        ("silenceMinutes", "60", "Wyciszenie beepów po 1×."),
        ("pushRepeatMinSeconds", "60", "Minimalny odstęp dla przypomnień push (nie eskaluje do 10 s jak syreny, żeby nie spamować)."),
        ("pumpRestoreDryMinutes", "5", "Brak WET przez X minut przed restore."),
        ("pumpRestoreDryReportMinutes", "10", "Wymóg świeżego DRY raportu z czujników."),
        ("repeatMinSeconds", "10", "Minimalna przerwa między powtórzeniami (eskalacja do tej wartości)."),
        ("floodRepeatSeconds", "150", "Startowa przerwa powtórzeń dla zalania (każde powtórzenie dzieli przez 2 aż do minimum)."),
        ("smokeRepeatSeconds", "120", "Startowa przerwa powtórzeń dla dymu (każde powtórzenie dzieli przez 2 aż do minimum)."),
        ("heatThresholdC", "45°C", "Próg alarmu temperatury."),
    ]:
        r = t.add_row().cells
        r[0].text = k
        r[1].text = str(v)
        r[2].text = d

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_PATH))
    print(f"Wrote {OUT_PATH}")


if __name__ == '__main__':
    main()
